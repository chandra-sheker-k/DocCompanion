import tempfile
from pathlib import Path
from unittest.mock import patch

from django.test import SimpleTestCase
from docx import Document
from docx.oxml import OxmlElement

from apps.documents.services import docx_extractor
from apps.documents.services.models import PageContent


class DocxPaginationTests(SimpleTestCase):
    @patch("apps.documents.services.docx_extractor._extract_rendered_pages")
    def test_rendered_pages_are_preserved(self, rendered_pages):
        rendered_pages.return_value = [
            PageContent(page=1, text="First page"),
            PageContent(page=2, text="Second page"),
        ]

        result = docx_extractor.extract("sample.docx")

        self.assertEqual([page.page for page in result.pages], [1, 2])
        self.assertEqual(result.metadata["page_number_source"], "rendered_layout")

    @patch("apps.documents.services.docx_extractor._extract_rendered_pages", return_value=None)
    def test_word_rendered_breaks_create_separate_pages(self, _rendered_pages):
        with tempfile.TemporaryDirectory() as temp_dir:
            path = Path(temp_dir) / "pages.docx"
            document = Document()
            document.add_paragraph("First page content")
            break_paragraph = document.add_paragraph()
            break_paragraph.add_run()._r.append(
                OxmlElement("w:lastRenderedPageBreak")
            )
            break_paragraph.add_run("Second page content")
            document.save(path)

            result = docx_extractor.extract(path)

        self.assertEqual(len(result.pages), 2)
        self.assertIn("First page content", result.pages[0].text)
        self.assertIn("Second page content", result.pages[1].text)
